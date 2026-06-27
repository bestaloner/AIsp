#!/usr/bin/env python3
"""Generate AI Model Tuning Plan DOCX V2.2 — Adds Standard-Answer JSON to V2.0 base."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    if level == 1: hs.font.size = Pt(18)
    elif level == 2: hs.font.size = Pt(14)
    else: hs.font.size = Pt(12)

# ── Helpers ──
def add_para(text, bold=False, size=Pt(10.5), align=None, color=None, spacing_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size; run.font.name = 'Microsoft YaHei'; run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = spacing_after
    return p

def add_simple_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        for r in c.paragraphs:
            for run in r.runs: run.font.size=Pt(9); run.font.bold=True; run.font.name='Microsoft YaHei'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(val)
            for r in c.paragraphs:
                for run in r.runs: run.font.size=Pt(9); run.font.name='Microsoft YaHei'
    doc.add_paragraph()
    return table

def add_callout(text, color=RGBColor(0x25, 0x63, 0xEB)):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.size=Pt(9.5); run.font.name='Microsoft YaHei'; run.font.color.rgb=color; run.bold=True
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5)
    for run in p.runs: run.font.size=Pt(10); run.font.name='Microsoft YaHei'
    return p

def add_code_block(text):
    """Add a shaded code block paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_diagram_marker(text):
    """Add a diagram insertion marker"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = 'Microsoft YaHei'; run.bold = True
    run.font.color.rgb = RGBColor(0xF5, 0x7A, 0x00)
    return p

doc.add_paragraph(); doc.add_paragraph()

# ═══════════════ COVER ═══════════════
add_para('AI隐患识别模型调优方案', bold=True, size=Pt(26), color=RGBColor(0x1D,0x4E,0xD8), align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(8))
add_para('VisionGuard · AI 能力平台', bold=True, size=Pt(14), color=RGBColor(0x64,0x74,0x8B), align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(24))
add_para('12类施工安全隐患 · YOLO + VLM 双模型微调 · 数据集建设 · 判定标准 · 标准答案JSON · 隐患库 · 规范引用', size=Pt(10), color=RGBColor(0x94,0xA3,0xB8), align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(32))
meta = [
    f'文档版本：V2.2',
    f'编制日期：{datetime.date.today().strftime("%Y-%m-%d")}',
    '适用项目：建筑施工AI安全巡检 — 能力平台',
    '关联系统：VisionGuard 业务平台 / AI能力平台',
    '文档状态：DRAFT · 待评审',
    'V2.2 更新说明：新增标准答案JSON格式设计，升级VLM训练数据格式，细化验收标准'
]
for m in meta: add_para(m, size=Pt(10), color=RGBColor(0x64,0x74,0x8B), align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=Pt(2))
doc.add_page_break()

# ═══════════════ 一、项目背景与目标 ═══════════════
doc.add_heading('一、项目背景与目标', level=1)
doc.add_heading('1.1 现状分析', level=2)
add_para(
    '当前 AI 隐患识别系统采用 YOLO（YOLO26s-v2.1，mAP 87.3%）与 VLM（Qwen3.6-27B via vLLM）'
    '双引擎架构，覆盖了建筑施工场景中 12 类常见安全隐患的检测与分类。在实际运行中，系统已经具备了基本的识别能力'
    '——它确实能"大致"看出哪里有问题，但这种"大致"离真正的生产级交付还有一段距离。')
add_para(
    '从用户体验的角度来看，当前的短板主要体现在四个方面。第一是误报率偏高——系统有时会将正常场景误判为隐患，'
    '例如将普通衣物判定为安全绳缺失，让复核人员产生"狼来了"的疲劳感。第二是漏报风险——部分小目标、被遮挡的目标、'
    '逆光场景中的目标无法有效检出，这恰恰是最危险的"真问题没有被看到"。第三是边界案例处理能力弱——对"部分佩戴安全帽"'
    '"远处人员""重叠物体"等模糊场景的判断缺乏一致性。第四是场景泛化不足——当前模型在特定光照条件、拍摄角度、'
    '工地类型下的表现差异较大。此外，新版 12 类隐患中新增了 3 类涉及"人数识别"的场景（施工电梯、吊篮、起重机械安拆），'
    '对 AI 的精细计数能力提出了更高要求。')
add_para(
    '客户方对此的评价是："大方向能看出来，但再深究我们就不是很专业了。"因此，本次调优的核心思路不是换模型、换架构，'
    '而是：先定标准，再建数据，最后微调优化。')

doc.add_heading('1.2 调优目标', level=2)
add_para(
    '本次调优的总体目标是：通过明确判定规则与输出标准、建设高质量标注数据集（含标准答案JSON）、'
    '对 YOLO 和 VLM 进行针对性微调，将 12 类隐患的综合识别准确率从当前水平提升至可交付的生产级标准。')
add_para(
    '具体量化目标：YOLO 综合 mAP@0.5 从约 87% 提升至不低于 92%；召回率从约 85% 提升至不低于 93%；'
    '精确率从约 78% 提升至不低于 85%。VLM 复核准确率从约 82% 提升至不低于 90%。'
    '端到端综合准确率从约 75% 提升至不低于 88%。任一类别的指标退化不超过 2%。')

doc.add_heading('1.3 工作路径', level=2)
add_para(
    '整个调优工作按照"定规则 → 做数据 → 练模型 → 验证上线"的路径推进，分为六个阶段。'
    '下图展示了从方案确认到全量上线的完整工作路径、各阶段的核心工作内容和时间安排。')
add_diagram_marker('【📊 插入图 1：整体工作路径图 — 文件 diagrams/01_整体工作路径图.drawio】')
add_para(
    '第一阶段确定范围与判定标准，逐一明确 12 类隐患的定义、判定规则和法规依据；'
    '第二阶段定义输出形式，设计"AI 检测 + 隐患库匹配 + 规范条款引用"的三层证据链输出结构；'
    '第三阶段数据集建设，按标准收集原始图像数据，完成图像标注与标准答案JSON的配套编写，经质量复核后入库；'
    '第四阶段 YOLO 微调，在现有权重基础上用标注数据进行针对性训练；'
    '第五阶段 VLM 微调，采用 LoRA 低秩适配方式，以标准答案JSON为训练目标对 Qwen3.6-27B 进行轻量化微调；'
    '第六阶段评估验证与灰度上线。')

doc.add_page_break()

# ═══════════════ 二、调优范围 ═══════════════
doc.add_heading('二、调优范围', level=1)
doc.add_heading('2.1 概述', level=2)
add_para(
    '本次调优覆盖全部 12 类施工安全隐患。按照隐患的性质，分为物理设施类（5 项）和人员行为类（7 项）。'
    '物理设施类隐患关注"物"的不安全状态——临边防护栏杆缺失、基坑边缘违规堆载、吊装工具不合格、围挡破损等。'
    '人员行为类隐患关注"人"的不安全行为——未佩戴安全帽、未系安全绳、违规吸烟、超员搭载、私自操作机械等。')
add_para(
    '特别值得注意的是，本次 12 类中有 3 类涉及"计数"判断——施工电梯搭载人数（不超过 9 人）、吊篮作业人数（不超过 2 人）、'
    '起重机械安拆人数（按方案要求）。这些类别不仅需要 AI 检测到"有人"，还需要精确统计人数并与法规上限比对。'
    '这与传统的"有/没有"二元检测存在本质区别，对 AI 的精细识别和场景理解能力提出了更高要求，在判定标准、标注规范和 JSON 标准答案中需要特别处理。')

add_simple_table(
    ['#', '隐患名称', '类别', '检测类型', '典型场景描述', '当前主要问题'],
    [
        ['1', '安全帽佩戴识别', '行为', '有/无', '施工区域人员头部未佩戴或未正确佩戴安全帽', '帽子颜色与背景相近时误判；部分佩戴状态漏检'],
        ['2', '安全绳佩戴识别', '行为', '有/无', '高处作业人员未系挂安全带或安全绳', '安全绳细长难检测；与脚手架钢管、电缆等线条物体易混淆'],
        ['3', '临边防护缺失识别', '物理设施', '有/无', '楼层边缘、洞口、阳台等临边位置缺少防护栏杆', '复杂背景干扰严重；不同类型护栏识别一致性差'],
        ['4', '施工电梯搭载人数识别', '行为', '计数', '施工升降机轿厢内搭载人数超过9人（含司机）', '需在轿厢狭小空间内精确计数；遮挡和重叠导致漏计'],
        ['5', '工人私自开机识别', '行为', '有/无', '非持证人员擅自操作施工机械（塔吊、施工电梯等）', '需识别操作人员身份与机械运行状态'],
        ['6', '坑边堆载识别', '物理设施', '有/无', '基坑边缘1.5m范围内违规堆放土方、材料、机械', '荷载类型多样且边界模糊；缺乏距离参照物'],
        ['7', '违规抽烟识别', '行为', '有/无', '施工区域内人员吸烟或出现烟头火光', '小目标检测困难；白色细长物体与香烟混淆'],
        ['8', '散料吊装识别', '物理设施', '有/无', '使用非标料斗吊装散料，或将钢筋挂成"雨伞状"吊装', '需组合判断吊具类型、物料形态和捆扎状态'],
        ['9', '吊篮作业人数识别', '物理设施', '计数', '高处作业吊篮内操作人员超过2人', '高空小目标计数困难；吊篮内人员重叠遮挡'],
        ['10', '起重机械安拆人数识别', '行为', '计数', '塔吊/施工电梯安拆作业时现场人数不符合方案要求', '需区分安拆人员与旁观人员；场景复杂人员分散'],
        ['11', '出入口安全防护棚识别', '物理设施', '有/无', '通道口未搭设安全防护棚或棚体尺寸不达标', '防护棚形式多样；需判断棚体高度和双层结构'],
        ['12', '工地围挡封闭识别', '物理设施', '有/无', '工地外围围挡破损、缺失、未完全封闭', '围挡形式多样，标准不统一'],
    ]
)

doc.add_heading('2.2 不在本次调优范围的事项', level=2)
add_para('为了聚焦资源达成核心目标，以下事项不在本次调优范围内：')
add_bullet('新增隐患类型的模型训练——本次仅针对现有 12 类进行精度优化，不扩展新的类别。')
add_bullet('平台业务功能变更——不涉及 VisionGuard 业务平台的代码修改或新功能开发。')
add_bullet('模型架构级别的替换——不更换 YOLO 或 VLM 的基础模型架构，仅在现有架构基础上进行微调。')
add_bullet('实时推理性能优化——本次聚焦于识别精度的提升，推理速度优化作为独立专项另行安排。')

doc.add_page_break()

# ═══════════════ 三、判定标准 ═══════════════
doc.add_heading('三、判定标准', level=1)
doc.add_heading('3.1 通用判定规则', level=2)
add_para(
    '判定标准是整个调优工作的根基。无论是标注人员标注数据、编写 JSON 标准答案，还是 AI 工程师评估模型，'
    '都需要一套统一的、可操作的、有据可依的判断准则。每类隐患的判定结论分为三种：'
    '正例——图像或视频中明确存在该类隐患，且具备可辨识的视觉特征，标注人员有 90% 以上把握确认；'
    '负例——明确不存在该类隐患的正常场景；'
    '边界样本——无法明确判定是否存在隐患的模糊场景（图像模糊、遮挡严重、光线不足、距离过远等），'
    '标注为 -1 并备注原因，不直接参与训练，但作为难例分析的依据。'
    '标注人员在把握低于 90% 时，应直接标记为"边界"并备注原因，不得强行判定。')

doc.add_heading('3.2 逐类判定细则', level=2)
add_para('以下对每类隐患给出详细的判定标准，融合了视觉特征分析和客户法规文档的双重依据。')

# ── #1 安全帽佩戴识别 ──
doc.add_heading('#1 安全帽佩戴识别', level=3)
add_para(
    '安全帽是施工现场最基本、最普遍的个体防护装备。JGJ 184-2009 第 2.0.4 条明确规定：'
    '"进入施工现场人员必须佩戴安全帽。"本类隐患的判断逻辑是"找未戴帽子的人"，'
    '而不是"判断这个人是不是该戴帽子"——只要在施工区域内出现的人员，都应当佩戴。'
    '正例：人员头部区域未出现安全帽轮廓（无论颜色），包括完全未佩戴、手持安全帽未戴上、'
    '佩戴普通帽子（棒球帽、头巾）替代。负例：人员头部清晰可见安全帽，或安全帽放置在地面/'
    '工具台上（非佩戴状态）。边界案例：头部被完全遮挡、人员距离过远（头部像素<20px）、'
    '摘/戴安全帽的动作瞬间。')

# ── #2 安全绳佩戴识别 ──
doc.add_heading('#2 安全绳佩戴识别', level=3)
add_para(
    '安全绳（安全带）是高处作业人员的"生命线"。JGJ 184-2009 第 2.0.4 条和 JGJ/T 429-2018 '
    '第 5.1.6 条明确要求：2m 以上的悬空作业人员必须佩戴安全带并正确系挂。'
    '本类隐患的检测难点在于安全绳本身非常细长，在图像中只占很少的像素，'
    '且容易与脚手架钢管、电缆线、钢丝绳等线条物体混淆。'
    '正例：高处作业（2m 以上）人员躯干和腰部区域未见安全带绑带特征、'
    '安全绳未连接到可靠的挂点上、安全绳处于松弛状态未受力。'
    '负例：人员可见完整安全带绑带且安全绳连接至可靠的锚固点。'
    '边界案例：安全带被衣物或工具遮挡、人员处于高处作业与非高处作业的临界高度（约 2m 附近）、'
    '安全绳与背景钢管重叠难以分辨。')

# ── #3 临边防护缺失识别 ──
doc.add_heading('#3 临边防护缺失识别', level=3)
add_para(
    '临边防护缺失是施工现场最常见的高处坠落隐患之一。GB 55023-2022 和 JGJ 80-2016 '
    '对各类临边、洞口的防护措施做出了详细规定。简单来说，凡是人可能坠落的地方，都必须有防护栏杆。'
    '正例包括：楼层边缘、楼梯口、阳台边等临空一侧未设置防护栏杆，或防护栏杆不完整——缺少上横杆、'
    '下横杆，立杆间距超过 2 米，栏杆高度不足 1.2 米。各类洞口（楼梯口、电梯井口、预留洞口）'
    '未设置防护栏杆或盖板覆盖。按 JGJ 80-2016 §4.2.1 条：竖向洞口短边<500mm 时应封堵，'
    '≥500mm 时应设置不低于 1.2m 的防护栏杆。负例：临边位置可见完整防护栏杆（上下两道横杆+'
    '立杆+踢脚板齐全），或用密目式安全立网全封闭的外脚手架。')

# ── #4 施工电梯搭载人数识别 ──
doc.add_heading('#4 施工电梯搭载人数识别（计数类）', level=3)
add_para(
    '本类是典型的"计数"型隐患。根据 JGJ 33-2012 第 2.0.7 条和韶建电[2017]68 号文的规定，'
    '施工升降机承载人数不得超过 9 人（含司机）。这是一个硬性数值限制，AI 需要在轿厢狭小封闭空间内'
    '准确统计人数。正例：施工电梯轿厢内可见人数 ≥10 人。检测的关键在于：识别出轿厢区域，'
    '对区域内的人员进行逐一计数——即使部分人员被遮挡，只要可见头部或躯干即算一人。'
    '负例：轿厢内可见人数 ≤9 人且每人都有可辨识的独立轮廓。'
    '边界案例：轿厢内人员严重重叠（难以分辨是 9 人还是 10 人）、轿厢门正在关闭过程中人员数量不确定、'
    '部分人员可能处于轿厢门槛区域难以判定是否已进入。'
    '此类对标注和标准答案 JSON 提出了更高要求：标注时需要框出轿厢区域并对每个独立人员标注计数标记，'
    'JSON 中需记录检测到的人数（count）和法规上限（count_limit）。')

# ── #5 工人私自开机识别 ──
doc.add_heading('#5 工人私自开机识别', level=3)
add_para(
    '本类隐患针对的是非授权人员擅自操作施工机械的行为。《重大事故隐患判定标准(2024)》第四条明确规定：'
    '"建筑施工特种作业人员未取得有效特种作业人员操作资格证书上岗作业"应判定为重大事故隐患。'
    'JGJ 33-2012 第 2.0.7 条也要求操作人员不得擅自离开工作岗位或将机械交给其他无证人员操作。'
    '正例：在施工机械操作台区域（塔吊驾驶室、施工电梯操作位、物料提升机控制位）发现人员正在操作，'
    '但该人员未穿着特种作业人员标识服装，或明显不具备操作资格的特征（如未持证、非指定操作人员）。'
    '负例：持证操作人员在操作位正常作业，或机械处于停机关闭状态且无人操作。边界案例：机械处于'
    '怠速运行状态但操作位无人（可能是自动运行或被远程控制）、操作区域有人员但无法判断是否在执行操作动作。')

# ── #6 坑边堆载识别 ──
doc.add_heading('#6 坑边堆载识别', level=3)
add_para(
    '坑边堆载是基坑工程中最容易引发坍塌事故的隐患之一。根据 JGJ 311-2013 第 11.2.2 条的条文说明：'
    '"基坑周边 1.5m 范围内不宜堆载，3m 以内限制堆载，坑边严禁重型车辆通行。"'
    'JGJ 120-2012 第 8.1.5 条进一步强调："基坑周边施工材料、设施或车辆荷载严禁超过设计要求的地面荷载限值。"'
    '正例：基坑边缘 1.5m 范围内出现任何堆土、建筑材料或施工机械，堆土高度超过 1.5m 且位于坑边，'
    '挖土机械在坑底边线周边与开挖深度相等范围内堆载或行走。还包括基坑周边未做硬化/防渗处理，'
    '坑体渗水积水未及时疏导。负例：荷载堆放距离基坑边缘 >1.5m 且有挡土墙/支护结构防护，'
    '施工机具在指定安全平台作业。边界案例：无法准确判断距离（缺乏参照物），荷载量处于临界值。')

# ── #7 违规抽烟识别 ──
doc.add_heading('#7 违规抽烟识别', level=3)
add_para(
    'GB 50720-2011 第 6.4.5 条明确规定："施工现场严禁吸烟。"这条规定没有任何例外——'
    '只要在施工区域内出现吸烟行为就是违规。正例：施工区域内人员手部或嘴部出现点燃的香烟，'
    '特征为可见烟头火光（红色亮点）或上升的烟雾。负例：手中持有其他细长物体但没有火光'
    '（如笔、小工具、钢筋头），或完全没有吸烟迹象。边界案例：疑似香烟但无法分辨（距离远/像素低），'
    '白色细长物体在手部附近但不能确认，电子烟（无明显火光特征）。')

# ── #8 散料吊装识别 ──
doc.add_heading('#8 散料吊装识别', level=3)
add_para(
    '闽建办建函[2024]60 号文件明确规定：吊装易散落材料时必须使用料斗吊运，严禁使用承插型盘扣式钢管等材料'
    '简易搭设非标料斗进行吊装作业，严禁将钢筋挂成"雨伞状"直接吊装。正例：使用钢管临时拼凑的料斗吊运散料、'
    '将钢筋挂成雨伞状直接吊装、吊运易散落材料时未使用标准料斗、料斗装载明显过满。'
    '关键在于识别吊具的类型——是标准料斗还是非标临时拼装物。负例：使用标准料斗吊运散料且装载量合规，'
    '钢筋绑扎成捆使用专用吊具吊装。边界案例：吊具类型无法从图像中清晰辨认，装载量难以准确判断。')

# ── #9 吊篮作业人数识别 ──
doc.add_heading('#9 吊篮作业人数识别（计数类）', level=3)
add_para(
    'JGJ 202-2010 第 5.5.8 条明确规定"吊篮内的作业人员不应超过 2 人"。JB/T 11699-2013 '
    '进一步规定双动力吊篮操作人员不允许单独一人。因此，吊篮内的合规人数区间为：双动力吊篮 2 人，'
    '单动力吊篮 1-2 人。正例：吊篮内可见 ≥3 人，或双动力吊篮仅见 1 人。'
    'AI 需要在高空、远距离、小目标的条件下对吊篮内人员进行精确计数，同时识别吊篮的动力类型。'
    '负例：吊篮内 ≤2 人且均佩戴安全带，双动力吊篮内恰为 2 人。'
    '边界案例：吊篮内人数难以准确计数（遮挡/角度/距离），安全绳/带是否固定无法确认。')

# ── #10 起重机械安拆人数识别 ──
doc.add_heading('#10 起重机械安拆人数识别（计数类）', level=3)
add_para(
    '塔吊和施工电梯的安装拆卸是施工现场风险最高的作业之一，必须严格按照专项施工方案执行，'
    '现场作业人数需要符合方案要求（通常为 4-6 人），且必须有持证的特种作业人员在场指挥。'
    '本类隐患的难点在于：安拆作业现场通常人员分散、场景复杂，AI 需要区分安拆作业人员与旁观人员，'
    '并对安拆人员进行计数。正例：安拆作业现场可见的实际作业人数明显少于或超过方案规定人数，'
    '或无持证人员在场。负例：安拆作业现场人数与方案一致且持证人员在岗。'
    '边界案例：部分人员处于作业区域边缘位置无法判断是否参与安拆作业、人员频繁移动进出导致计数不稳定。')

# ── #11 出入口安全防护棚识别 ──
doc.add_heading('#11 出入口安全防护棚识别', level=3)
add_para(
    'JGJ 80-2016 第 7.2.1 条和 JGJ/T 429-2018 第 6.0.1、6.0.2 条对安全防护棚做出了详细规定：'
    '人员进出的通道口（包括物料提升机、施工升降机的进出通道口）必须设置安全防护棚；'
    '非机动车通行时棚底至地面高度不应小于 3m，机动车通行时不应小于 4m；'
    '当建筑物高度大于 24m 且采用木质板搭设时，应搭设双层安全防护棚，两层间距不应小于 700mm。'
    '正例：通道口完全未搭设防护棚，或棚体尺寸明显不达标（高度不足、单层未达双层要求）。'
    '负例：通道口上方可见符合标准的完整安全防护棚。边界案例：无法判断棚体实际高度和层数。')

# ── #12 工地围挡封闭识别 ──
doc.add_heading('#12 工地围挡封闭识别', level=3)
add_para(
    'JGJ 59-2011 第 3.2.3 条规定：市区主要路段的工地应设置高度不小于 2.5m 的封闭围挡，'
    '一般路段的工地应设置高度不小于 1.8m 的封闭围挡，围挡应坚固、稳定、整洁、美观。'
    '正例：围挡高度明显不达标、围挡出现破损或倒塌、围挡未连续封闭存在缺口、围挡明显倾斜。'
    '负例：围挡高度达标且连续封闭、结构稳固、外观整洁。边界案例：无法判断围挡实际高度（缺乏参照物），'
    '围挡部分被遮挡无法确认完整性。')

doc.add_page_break()

# ═══════════════ 三-B. 法规规范依据 ═══════════════
doc.add_heading('3.3 法规规范依据', level=2)
add_para(
    '以下法规依据来源于客户方提供的《远程视频监控常见隐患场景+定义》PDF 文档。'
    '12 类隐患与 PDF 的覆盖匹配度——PDF 中的 11 项规范定义中有 9 项直接对应到新的 12 类中。'
    '这些法规条款也将在标准答案 JSON 的 regulation_clause 和 regulation_text 字段中作为预填内容使用。')
add_para(
    '已有法规依据的 9 类：安全帽佩戴识别对应 JGJ 184-2009 §2.0.4；安全绳佩戴识别对应 JGJ 184-2009 §2.0.4 '
    '和 JGJ/T 429-2018 §5.1.6；临边防护缺失识别对应 GB 55023-2022 §4.4.4/§5.2.1 和 JGJ 80-2016 '
    '§4.1.1-§4.2.1；施工电梯搭载人数识别对应 JGJ 33-2012 §2.0.7 和韶建电[2017]68 号（限载 9 人）；'
    '坑边堆载识别对应 JGJ 120-2012 §8.1.5/§8.1.6、JGJ 311-2013 §8.3.2/§11.2.2 和 GB 55034-2022 §3.5.3 等多项；'
    '违规抽烟识别对应 GB 50720-2011 §6.4.5；散料吊装识别对应闽建办建函[2024]60 号；'
    '吊篮作业人数识别对应 JGJ 202-2010 §5.5.8（≤2 人）和 JB/T 11699-2013；'
    '出入口安全防护棚识别对应 JGJ 80-2016 §7.2.1 和 JGJ/T 429-2018 §6.0.1/§6.0.2。')
add_para(
    '另有 3 类暂无客户提供的法规定义文档——工人私自开机识别（#5）、起重机械安拆人数识别（#10）、'
    '工地围挡封闭识别（#12，JGJ 59-2011 有规定但 PDF 中提取不够完整）。'
    '建议标注工作启动前，与客户方确认这 3 类的法规依据，或补充相关规范文档。'
    '确认后同步更新标准答案 JSON 中对应的预填字段。')

doc.add_page_break()

# ═══════════════ 四、输出形式定义 ═══════════════
doc.add_heading('四、输出形式定义', level=1)
doc.add_heading('4.1 设计理念：双重依据架构', level=2)
add_para(
    '传统的 AI 隐患检测输出只回答了两个问题："这是什么隐患"和"在哪"。施工现场的安全管理'
    '需要回答第三个也是最重要的问题："凭什么说它是隐患"。这个"凭什么"有两个来源：'
    '一是隐患库（Hazard Knowledge Base）——从知识定义层面说明为什么这个场景被判定为隐患；'
    '二是规范条款——从法规合规层面说明违反了哪条国家或行业标准，原文是什么。'
    '下图展示了从摄像头抓拍到完整证据链输出的三层架构流程。')
add_diagram_marker('【📊 插入图 2：证据链输出流程图 — 文件 diagrams/02_证据链输出流程图.drawio】')
add_para(
    '输出格式设计为三层证据链。最底层是 AI 检测结果（YOLO + VLM 从图像中识别到的目标信息和置信度）；'
    '中间层是隐患库匹配（从结构化隐患知识库中查找匹配的条目，展开完整的定义、视觉特征和严重级别）；'
    '最上层是规范条款引用（从法规文档中提取的 GB/JGJ 标准条款，包括编号和原文）。'
    '对复核人员来说，看到的不再只是一个冷冰冰的标签，而是"这是什么—为什么是—违反了什么"的完整逻辑链。')

doc.add_heading('4.2 模型检测输出格式', level=2)
add_para(
    '单次检测输出 JSON 格式，每个隐患对象包含三个区块。基础检测信息：检测 ID、时间戳、摄像头 ID、'
    '项目 ID、隐患类型编号、边界框坐标、YOLO 置信度、VLM 复核结果子对象。隐患库匹配信息（本次新增）：'
    '隐患库条目 ID、隐患标准名称、隐患定义、AI 判断所依据的视觉特征列表、严重级别、匹配置信度。'
    '规范条款引用（本次新增）：法规条款 ID、法规全称、条款号、条款原文、违规事实描述。'
    '此外，对于 3 类计数型隐患（#4/#9/#10），输出中增加 count 字段，'
    '记录 AI 检测到的实际人数并与法规上限进行比对。')
add_callout(
    '证据链示例：AI检测[安全绳未佩戴,置信度0.91] → 隐患库匹配[KB-HZ-002, 安全绳佩戴识别, '
    '定义：高处作业人员未系挂安全带或安全绳] → 违反[JGJ 184-2009 §2.0.4: '
    '"2m及以上的无可靠安全防护设施的高处作业时，必须系挂安全带"]')

doc.add_heading('4.3 隐患库（Hazard Knowledge Base）设计', level=2)
add_para(
    '隐患库是连接"AI 看到了什么"和"这件事意味着什么"的核心纽带。每条隐患记录是一个结构化数据条目，包含：'
    '唯一编号（KB-HZ-XXX）、对应的 class_id、中英文名称、分类、定义、AI 检测视觉特征列表、'
    '正例/负例/边界案例的文字描述、默认严重级别、以及关联法规列表。'
    '隐患库以 JSON 文件形式维护（hazard_knowledge_base.json），存放于数据集 metadata 目录下。'
    '当 AI 检测到一个隐患时，系统根据 class_id 自动检索对应的 KB 条目，将结构化定义和法规引用附加到输出中。')

# ── NEW: 4.4 标准答案JSON：训练与输出的统一格式 ──
doc.add_heading('4.4 标准答案JSON：训练与输出的统一格式', level=2)
add_para(
    '本次调优的核心设计思想之一是"训练即输出"——VLM 在训练阶段学习的 JSON 标准答案格式，'
    '与推理阶段输出的 JSON 格式保持一致。这确保了模型学会的恰好是它需要产出的，'
    '避免了传统方案中"先训练分类标签、再靠后处理拼接证据链"的信息损失。')
add_para(
    '训练 JSON（标准答案）与推理输出 JSON 的对应关系如下。二者结构高度一致，区别仅在于字段值的来源：'
    '训练 JSON 中所有字段均由人工标注和隐患库预填产生（是"标准答案"），'
    '推理输出 JSON 中同名字段由 AI 推断产生。')

add_simple_table(
    ['字段组', '训练JSON（人工标注）', '推理输出JSON（AI推断）', '说明'],
    [
        ['image / image_info', '标注工具自动读取', '系统自动填充', '图像元信息，格式完全一致'],
        ['annotations[].bbox', '人工绘制边界框', 'YOLO 预测边界框', '坐标格式相同，来源不同'],
        ['annotations[].class_id', '人工选择类别', 'YOLO + VLM 联合判定', '类别编号体系相同'],
        ['annotations[].count', '人工计数填写', 'AI 统计算法输出', '仅计数型3类有此字段'],
        ['hazard_content.inspection_item', '隐患库预填，人工确认', 'VLM 推理生成', '取值来自同一隐患库'],
        ['hazard_content.violation_description', '人工根据模板撰写', 'VLM 推理生成', '自由文本，训练时学习表达模式'],
        ['hazard_content.regulation_clause', '隐患库预填，人工确认', 'VLM 推理生成 / KB检索', '静态字段，预填后仅需确认'],
        ['hazard_content.regulation_text', '隐患库预填，人工确认', 'VLM 推理生成 / KB检索', '静态字段，预填后仅需确认'],
    ]
)

add_para(
    '这种统一格式设计的优势在于：VLM 学习的不只是"图像中有没有安全绳"，而是从图像到完整证据链的端到端映射。'
    '模型在训练过程中反复接触"图像→检出框+类别+检查项目+违规描述+法规条款+法规原文"的完整示例，'
    '在推理时就能直接产出同样结构的输出，无需额外的后处理拼接步骤。')

doc.add_page_break()

# ═══════════════ 五、数据集建设方案 ═══════════════
doc.add_heading('五、数据集建设方案', level=1)
doc.add_heading('5.1 数据采集要求', level=2)
add_para(
    '高质量的数据是模型微调成功的前提。下图展示了从原始数据采集到标注完成、再到模型训练的全链路技术流程。')
add_diagram_marker('【📊 插入图 3：标注与微调技术流程图 — 文件 diagrams/03_标注与微调技术流程图.drawio】')
add_para(
    '数据来源优先使用现有项目中已部署摄像头的历史抓拍数据——'
    '这些数据真实反映了模型在实际使用中遇到的场景。如果现有数据在某些类别、某些场景上数量不足，'
    '则需要组织专项补充拍摄来补齐短板。')
add_para(
    '场景覆盖要求：至少 3 个不同施工阶段（基础/主体/装修）的不同工地；光照需覆盖白天、阴天、黄昏、'
    '夜间（含补光）和逆光场景，每种占比不低于 10%；天气需覆盖晴天、阴天、雨天，每种占比不低于 10%；'
    '拍摄角度需包括俯视、平视和仰视画面。图像格式统一为 JPG/PNG，分辨率不低于 1280×720，建议 1920×1080。')
add_para(
    '对于 3 类计数型隐患，数据采集时需要特别注意：每类至少 200 张含不同人数配置的图像（如施工电梯 4-12 人的场景），'
    '确保模型学习到人数从合规到违规的渐变过程，而不是只见过合规或只见过严重违规的极端情况。')
add_para(
    '本次数据集建设的核心升级点在于：除了传统的图像+边界框标注外，每张训练图像将配套生成一份 JSON 格式的'
    '标准答案文件，包含人工标记、检查项目、违规描述、规定条款和规定原文，作为 AI 训练的完整"参考答案"。'
    '有关标准答案 JSON 的详细格式设计，见下一节 5.2。')

# ── NEW: 5.2 标准答案JSON格式设计 ──
doc.add_heading('5.2 标准答案JSON格式设计', level=2)
add_para(
    '标准答案 JSON 是整个数据集建设的核心交付物之一。它将传统的"图像+标注框"升级为'
    '"图像+标注框+结构化证据链"，让 AI 在训练时就能学习到从视觉特征到法规判定的完整推理过程。')

doc.add_heading('5.2.1 JSON 顶层结构', level=3)
add_para(
    '每条标准答案 JSON 对应一张训练图像。顶层包含三个一级字段：image（图像文件信息）、'
    'image_info（拍摄元信息）、annotations（标注对象数组，每个元素描述图像中的一个隐患实例）。'
    '当一张图像包含多个隐患时，annotations 数组中包含多个元素。'
    '以下为一个完整的标准答案 JSON 示例（安全绳缺失场景）：')

# JSON 示例
json_example_1 = '''{
  "dataset_version": "1.0",
  "image": {
    "file_name": "CAM03_20250601_143000.jpg",
    "width": 1920,
    "height": 1080
  },
  "image_info": {
    "camera_id": "CAM-03",
    "capture_time": "2025-06-01T14:30:00",
    "light_condition": "daylight",
    "weather": "sunny",
    "construction_phase": "main_structure"
  },
  "annotations": [
    {
      "annotation_id": 1,
      "bbox": {"x": 320, "y": 180, "width": 120, "height": 280},
      "class_id": 1,
      "class_name": "安全绳佩戴识别",
      "detection_type": "有/无",
      "is_positive": true,
      "is_borderline": false,
      "borderline_reason": "",
      "count": null,
      "count_limit": null,
      "annotator_id": "A001",
      "annotation_date": "2025-07-03",
      "hazard_content": {
        "inspection_item": "安全绳佩戴识别",
        "violation_description": "2层外脚手架上一名作业人员腰部未见安全带绑带特征，安全绳未连接至可靠挂点",
        "regulation_clause": "JGJ 184-2009 §2.0.4",
        "regulation_text": "在2m及以上的无可靠安全防护设施的高处作业时，必须系挂安全带。",
        "kb_entry_id": "KB-HZ-002",
        "severity": "高"
      }
    }
  ]
}'''
add_code_block(json_example_1)

doc.add_heading('5.2.2 字段详细说明', level=3)
add_para(
    '以下逐字段说明每个字段的含义、数据类型、是否必填和填写方式。理解这些字段的用途，'
    '有助于标注人员准确填写，也有助于 AI 工程师正确解析训练数据。')

add_para('【顶层字段】', bold=True, size=Pt(10.5))
add_para(
    'dataset_version：数据集版本号，字符串，必填。由数据集管理员在创建批次时统一赋值，用于追溯数据版本。')
add_para(
    'image 对象：包含 file_name（图像文件名，字符串，必填，标注工具自动读取）、width 和 height（图像像素尺寸，整数，必填，自动读取）。')
add_para(
    'image_info 对象：包含 camera_id（摄像头编号，字符串，必填）、capture_time（抓拍时间戳，ISO 8601 格式，必填）、'
    'light_condition（光照条件，枚举值：daylight/overcast/dusk/night/backlight，必填）、'
    'weather（天气，枚举值：sunny/cloudy/rainy，必填）、construction_phase（施工阶段，枚举值：foundation/main_structure/decoration，必填）。'
    'image_info 中的场景维度字段用于后续按维度统计数据集分布，确保训练数据覆盖各类场景。')

add_para('【annotations 数组 — 每个隐患实例的字段】', bold=True, size=Pt(10.5))
add_para(
    'annotation_id：标注序号，整数，必填。在同一张图像的 annotations 数组中从 1 开始递增。')
add_para(
    'bbox 对象：包含 x、y（边界框左上角坐标，整数，必填）、width、height（边界框宽高，整数，必填）。'
    '坐标原点为图像左上角。标注人员使用 LabelImg/Label Studio 等工具绘制时自动生成。')
add_para(
    'class_id：隐患类别编号，整数 0-11，必填。对应附录 A 中的 class_id 映射表。标注人员选择类别标签后自动填入。')
add_para(
    'class_name：隐患中文名称，字符串，必填。根据 class_id 自动填入。')
add_para(
    'detection_type：检测类型，"有/无"或"计数"，字符串，必填。根据 class_id 自动填入。')
add_para(
    'is_positive：是否为正例，布尔值，必填。true=正例（存在隐患），false=负例（不存在隐患）。')
add_para(
    'is_borderline：是否为边界样本，布尔值，必填。若为 true，该条目不参与训练，仅用于难例分析。'
    '若为 true，borderline_reason 字段必须填写具体原因。')
add_para(
    'borderline_reason：边界样本原因，字符串。当 is_borderline=true 时必填，否则留空。'
    '常见原因包括：图像模糊、遮挡严重、光线不足、距离过远、目标处于临界状态。')
add_para(
    'count：AI 检测到的人数，整数。仅 detection_type="计数"时填写，其他类型填 null。'
    '标注人员根据实际可见人数手动填写，用于训练 VLM 的计数能力。')
add_para(
    'count_limit：法规人数上限，整数。仅 detection_type="计数"时填写，其他类型填 null。'
    '根据法规自动预填（如 #4 施工电梯填 9，#9 吊篮填 2），标注人员确认。')
add_para(
    'annotator_id：标注人员编号，字符串，必填。用于追溯标注来源和按人员统计质量。')
add_para(
    'annotation_date：标注日期，YYYY-MM-DD 格式，字符串，必填。')

add_para('【hazard_content 对象 — 隐患内容（标准答案核心字段）】', bold=True, size=Pt(10.5))
add_para(
    'inspection_item：检查项目名称，字符串，必填。与 class_name 相同，由隐患库预填，标注人员确认即可。')
add_para(
    'violation_description：违规描述，字符串，必填。用 1-3 句自然语言具体描述图像中发现的违规事实。'
    '这是整个标准答案 JSON 中唯一需要标注人员手工撰写的文本字段。为保持一致性，'
    '每类隐患提供 3-5 个标准描述模板（见 5.2.3 节），标注人员从中选择最匹配的进行微调。'
    '描述要点：指明具体位置（如"2层外脚手架"）、指出缺失/异常的对象（如"腰部未见安全带绑带"）、'
    '说明与标准的差距（如"安全绳未连接至可靠挂点"）。')
add_para(
    'regulation_clause：规定条款编号，字符串，必填。由隐患库根据 class_id 自动预填，标注人员确认。'
    '若图像中违规事实对应多条法规，以最主要的一条为准。特殊情况下标注人员可手动切换为其他适用条款。')
add_para(
    'regulation_text：规定原文，字符串，必填。由隐患库根据 class_id 自动预填，标注人员确认。'
    '内容为法规条款的完整原文引用。')
add_para(
    'kb_entry_id：隐患库条目编号，字符串，必填。由隐患库根据 class_id 自动预填。'
    '用于将标准答案与隐患知识库条目关联，确保术语和定义的一致性。')
add_para(
    'severity：严重级别，枚举值：高/中/低，字符串，必填。由隐患库预填默认值，'
    '标注人员可根据具体违规的严重程度手动调整（如未佩戴安全绳且作业高度>10m 可调高为"高"）。')

doc.add_heading('5.2.3 静态字段预填机制', level=3)
add_para(
    '为减少标注人员的工作量并保证数据一致性，hazard_content 中的 inspection_item、regulation_clause、'
    'regulation_text、kb_entry_id、severity 这 5 个字段在标注工具中根据 class_id 自动从隐患库预填。'
    '标注人员只需在标注工具中确认预填内容是否正确，无需手工输入。仅当预填内容与实际情况不符时'
    '（如同一类别在不同场景下适用不同法规条款），才需要手动切换或修改。')
add_para(
    'violation_description 是唯一需要标注人员手工撰写的字段。为降低撰写难度并保持风格一致性，'
    '每类隐患提供 3-5 个标准描述模板作为起点。以 #2 安全绳佩戴识别为例，三个模板分别是：'
    '"（位置）作业人员腰部未见安全带绑带特征，安全绳未连接至可靠挂点"、'
    '"（位置）作业人员虽佩戴安全带但安全绳处于松弛状态，未有效系挂"、'
    '"（位置）作业人员安全绳挂点不可靠/不符合规范要求"。'
    '标注人员根据图像中实际违规表现选择最匹配的模板，修改位置信息和具体细节后填入。')
add_para(
    '这种"预填+确认+模板辅助"的设计，将每张图像的标准答案 JSON 填写时间控制在 2-3 分钟以内，'
    '不会给标注流程带来过重的额外负担。')

doc.add_heading('5.2.4 多隐患图像的处理', level=3)
add_para(
    '施工场景中一张图像往往同时包含多种隐患（如同一画面中既有安全帽缺失、又有临边防护缺失）。'
    '标准答案 JSON 通过 annotations 数组天然支持这种多隐患场景——每个隐患实例作为数组中的一个独立元素，'
    '各自携带完整的 bbox、class_id 和 hazard_content。')
add_para(
    '对于多隐患图像，标注流程不变：标注人员依次绘制每个隐患的边界框、选择类别，'
    '每个隐患的 hazard_content 独立填写。Annotations 数组的长度等于该图像中隐患实例的总数。'
    '一套标注工具需支持在同一图像上添加、编辑、删除多个标注对象，并分别填写各自的隐患内容。')

doc.add_heading('5.2.5 计数型类别的特殊字段', level=3)
add_para(
    '3 类计数型隐患（#4 施工电梯搭载人数、#9 吊篮作业人数、#10 起重机械安拆人数）'
    '在 annotations 中额外包含两个字段：')
add_bullet('count：AI 检测到的实际人数（整数）。标注人员需在图像中逐一计数可见人员后填写。对于轿厢/吊篮内的人员，即使部分遮挡，只要可见头部或躯干即算一人。')
add_bullet('count_limit：法规规定的人数上限（整数）。由隐患库预填（#4 = 9，#9 = 2，#10 = 按方案），标注人员确认。')
add_para(
    '对于正例（存在违规），count 应大于 count_limit。对于负例（合规场景），count 应 ≤ count_limit。'
    '对于边界样本（人数恰好处于临界值附近），is_borderline 标记为 true 并在 borderline_reason 中说明原因。'
    'VLM 在训练时将同时学习"判定是否违规"和"比对计数结果"，使其在推理时能够直接输出带有计数比对的结构化结论。')

# 计数型 JSON 示例
json_example_count = '''{
  "annotations": [
    {
      "annotation_id": 1,
      "bbox": {"x": 120, "y": 200, "width": 350, "height": 420},
      "class_id": 3,
      "class_name": "施工电梯搭载人数识别",
      "detection_type": "计数",
      "is_positive": true,
      "is_borderline": false,
      "count": 11,
      "count_limit": 9,
      "annotator_id": "A003",
      "annotation_date": "2025-07-05",
      "hazard_content": {
        "inspection_item": "施工电梯搭载人数识别",
        "violation_description": "施工电梯轿厢内可见11人（含司机），超过法规上限9人",
        "regulation_clause": "JGJ 33-2012 §2.0.7; 韶建电[2017]68号",
        "regulation_text": "承载人数不得超过9人（含司机）。",
        "kb_entry_id": "KB-HZ-004",
        "severity": "高"
      }
    }
  ]
}'''
add_para('以下为计数型隐患的 JSON 示例（施工电梯超载场景）：', bold=True, size=Pt(10))
add_code_block(json_example_count)

doc.add_page_break()

# ── 5.3 数据量规划（原5.2重新编号） ──
doc.add_heading('5.3 数据量规划', level=2)
add_para(
    '每类正例不少于 500 张、负例不少于 100 张、边界样本不少于 50 张（不入训练集，仅备查）。'
    '12 类合计训练集不少于 6,000 张、验证集不少于 1,000 张、测试集不少于 1,000 张，'
    '总计标注样本不少于 8,520 个，加上约 600 个边界样本。计数型 3 类的正例中需额外包含不同人数配置的标注信息。'
    '独立图片总数约为标注样本数的 60%-70%（因一张图常含多类隐患），即约 5,000-6,000 张独立图片。')
add_para(
    '由于新增了标准答案 JSON 的填写工作（每张图像增加 2-3 分钟），标注效率从"每人日均 50-80 张"'
    '调整为"每人日均 40-60 张"。建议配置 4 名标注人员加 1 名组长，标注总周期约 4-5 周（较原估算增加 1-2 周）。'
    '如能通过模板预填和工具优化进一步减少文本填写时间，周期可压缩回 3-4 周。')

# ── 5.4 标注流程与质量控制（原5.3重新编号） ──
doc.add_heading('5.4 标注流程与质量控制', level=2)
add_para(
    '标注流程六步（较原方案增加标准答案编写步骤）：')
add_bullet('第一步 — 原始数据采集：从现有系统导出历史抓拍图像，按摄像头/日期归档。缺失场景组织专项补充拍摄。')
add_bullet('第二步 — 数据清洗：剔除模糊/过曝/纯黑图像，去除重复帧（相邻帧相似度>95%仅保留1张），确保可用率≥95%。')
add_bullet('第三步 — 图像标注：使用 LabelImg/Label Studio 按第三章判定标准绘制边界框 + 选择类别标签 + 标记正例/负例/边界。计数型类别额外标注人数。')
add_bullet('第四步 — 标准答案编写（新增）：在标注工具中确认预填的检查项目/规定条款/规定原文，根据模板撰写违规描述。计数型填写实际人数。')
add_bullet('第五步 — 质量复核：每批次随机抽 20% 交叉验证。边界框 IoU≥0.85，标签一致率≥95%。JSON 字段完整性 100%，违规描述与标注类别一致性≥95%。不一致率>10% 则整批返工。')
add_bullet('第六步 — 数据集入库：按 70/15/15 划分 train/val/test。生成类别分布/场景分布/光照分布统计报告。JSON 文件通过 Schema 验证后与图像一同入库。数据集版本冻结。')
add_para(
    '标准答案 JSON 的质量控制是本次新增的关键环节。除边界框验证外，复核人员需抽检 JSON 文本字段的质量，'
    '重点检查：违规描述是否与标注类别一致（如 #1 安全帽的违规描述不应写成安全绳相关内容）、'
    '法规条款引用是否正确、计数型类别的 count 和 count_limit 是否填写、'
    '必填字段是否有空值。JSON 格式校验采用自动化脚本，100% 文件通过 JSON Schema 验证后方可入库。')

doc.add_page_break()

# ═══════════════ 六、模型微调技术方案 ═══════════════
doc.add_heading('六、模型微调技术方案', level=1)
doc.add_heading('6.1 YOLO 微调方案', level=2)
add_para(
    'YOLO 微调采用"在现有权重基础上继续训练"的思路。当前 YOLO 版本已具备通用目标检测能力，'
    '微调要做的是让它把这些通用知识更精准地应用到施工安全这个特定领域。使用 Ultralytics YOLO 标准训练框架，'
    '图像尺寸保持 640×640，训练 100-150 个 epoch。微调阶段初始学习率设为 0.001，比从头训练低一个数量级，'
    '目的是让模型在现有知识基础上做小幅调整而非大幅度改变。学习率采用余弦调度逐渐衰减，'
    '配合早停机制（连续 15 epoch 验证集 mAP 无提升即自动停止）防过拟合。数据增强采用 Mosaic+MixUp+HSV 抖动+随机翻转组合。'
    '输出权重命名为 YOLO26s-v2.2-ft。')
add_para(
    '对于计数型类别（#4/#9/#10），微调时还需要关注模型对密集小目标的检测能力。'
    '考虑在训练集中适当增加多人员场景的占比，并评估是否需要引入专门的计数损失函数或密度图辅助分支。'
    'YOLO 训练数据由标准答案 JSON 中的 bbox 和 class_id 字段导出为 .txt 格式，不涉及 JSON 文本字段的使用。')

doc.add_heading('6.2 VLM 微调方案', level=2)
add_para(
    'VLM 微调采用 LoRA（Low-Rank Adaptation）技术。通俗地说，不是在修改这个 270 亿参数的大模型本身，'
    '而是在它的关键部位"贴几张便利贴"——这些便利贴包含专门针对施工安全场景的判断知识，体积仅几十 MB。'
    '使用时基础模型加 LoRA 适配器即可输出更专业的判断；不需要时可随时卸载恢复原始模型。')
add_para(
    '本次 VLM 微调的核心升级在于训练数据格式：由原来的"图像+QA 对"升级为"图像+标准答案 JSON"。'
    '训练前，标准答案 JSON 通过预处理脚本转换为 VLM 可消费的指令微调格式——'
    '将 image 作为视觉输入，将 hazard_content 中的 inspection_item、violation_description、'
    'regulation_clause、regulation_text 等字段按固定模板拼接为结构化期望输出。'
    '指令模板示例：给定图像，请检测施工安全隐患，输出 JSON 格式结果，包含检查项目、违规描述、规定条款和规定原文。'
    '这种"训练即输出"的设计让 VLM 直接学会输出完整的证据链 JSON，而非先输出类别标签再靠后处理拼接。')
add_para(
    'LoRA 参数配置：rank=16，alpha=32，训练数据每类≥200 条标准答案 JSON（3 类计数型各≥250 条），'
    '训练 3-5 个 epoch，学习率 2e-5，优化器 AdamW。'
    '微调目标是提升 VLM 在施工安全特定场景下的判断准确性，同时让其学会输出符合 Schema 约束的结构化 JSON。')

doc.add_heading('6.3 评估与上线', level=2)
add_para(
    '离线评估：测试集上计算全部指标，确保达标且无类别退化超 2%。同时评估 VLM 输出的 JSON 格式合规率'
    '（能通过 Schema 验证的比例）。难例专项：边界样本集测试，误判率降低≥20%。'
    'A/B 灰度：选取 2-3 个典型摄像头双路推理，人工盲评比对 1 周。灰度准确率提升≥10% 则全量切换，'
    '持续监控 2 周确认稳定后正式结项。')

doc.add_page_break()

# ═══════════════ 七、角色分工与协作流程 ═══════════════
doc.add_heading('七、角色分工与协作流程', level=1)
doc.add_heading('7.1 角色职责', level=2)
add_para('项目负责人（双方各一人）：整体进度把控、跨团队协调、关键决策确认和争议仲裁。')
add_para('数据采集人员（客户方）：从现有系统导出历史抓拍图像，按场景类别/光照/天气等维度分类归档；缺失维度组织补充拍摄。')
add_para(
    '标注人员（客户方为主，4人，我方提供培训和技术支持）：按判定标准和标注手册逐张绘制边界框、选择类别标签、'
    '填写标准答案 JSON 的文本字段。需经过法规知识基础培训（半天），能够根据图像中的违规事实选择匹配的违规描述模板并进行微调。'
    '如客户方标注人员不具备法规知识，由我方提供完整的描述模板库，标注人员从中选择最匹配的并确认。')
add_para(
    '标注复核员（我方）：每批次随机抽 20% 交叉验证。除边界框质量检查外，新增 JSON 文本字段质量抽检——'
    '违规描述与标注类别一致性检查、法规条款引用正确性验证、计数型字段完整性确认。'
    '对不一致案例做最终判定，管理数据集版本。')
add_para(
    '知识库维护员（我方，新增角色）：在标注启动前完成 12 类隐患的隐患库 JSON 和法规条款库的静态字段配置，'
    '确保标注工具中预填的 inspection_item、regulation_clause、regulation_text、kb_entry_id、'
    'severity 字段准确无误。标注过程中根据反馈持续维护和更新预填内容。')
add_para('AI 工程师（我方）：数据预处理、标准答案 JSON 到训练格式的转换脚本开发、YOLO 微调训练、VLM LoRA 微调训练、离线评估、灰度上线技术支持。')
add_para('业务验收人（客户方）：在灰度测试阶段验证新模型实际表现，包括验证 VLM 输出的证据链 JSON 是否可用、违规描述是否合理，给出最终验收意见。')

doc.add_heading('7.2 协作流程', level=2)
add_para(
    '下图展示了客户方与我方在调优各阶段的角色分工、交付物流转关系和沟通渠道。')
add_diagram_marker('【📊 插入图 4：协作角色图 — 文件 diagrams/04_协作角色图.drawio】')
add_para(
    'Phase 1 — 方案确认（第 1 周）：评审并签字确认方案和判定标准，确认 12 类隐患的描述模板库，'
    '编排《标注指导手册》（含标准答案 JSON 字段填写指南），部署标注工具（含 JSON 预填功能）。')
add_para(
    'Phase 2 — 数据采集与清洗（第 2-3 周）：客户采集交付原始数据，我方执行清洗去重，输出清洗报告。'
    '知识库维护员同步完成隐患库预填字段的配置和验证。')
add_para(
    'Phase 3 — 数据标注与复核（第 4-8 周，含标准答案 JSON 填写）：标注人员按标准标注图像边界框并填写 JSON 标准答案，'
    '每批次完成后复核人员交叉验证（边界框+JSON 文本字段双维度）。')
add_para(
    'Phase 4 — 模型微调训练（第 9-10 周）：JSON 标准答案数据经预处理脚本转换为训练格式，'
    'YOLO + VLM 双模型微调，离线评估（含 JSON Schema 合规率评估）。')
add_para(
    'Phase 5 — 验证与上线（第 11-12 周）：A/B 灰度双路测试→评估确认→全量切换。总周期约 12 周（较原方案增加约 2 周，'
    '主要增加在标注阶段的 JSON 填写时间）。')

add_simple_table(
    ['阶段', '工作内容', '周期', '关键交付物'],
    [
        ['P1 方案确认', '评审方案、编写标注手册（含JSON指南）、部署标注工具', '第 1 周', '签字方案 + 标注指导手册 + 描述模板库 + 工具就绪'],
        ['P2 数据采集', '导出历史数据、补充拍摄、清洗去重、KB预填配置', '第 2-3 周', '原始数据集 + 数据清洗报告 + 隐患库预填配置'],
        ['P3 数据标注', '图像标注 + JSON标准答案编写 + 双维度交叉验证', '第 4-8 周', '标注完成数据集 + 标准答案JSON + 质量报告'],
        ['P4 模型微调', 'JSON转训练格式 + YOLO微调 + VLM LoRA微调 + 评估', '第 9-10 周', '微调权重 + 评估对比报告 + Schema合规率报告'],
        ['P5 验证上线', 'A/B灰度双路测试 → 全量切换', '第 11-12 周', '灰度报告 + 上线确认 + 项目总结'],
    ]
)

doc.add_page_break()

# ═══════════════ 八、验收标准 ═══════════════
doc.add_heading('八、验收标准', level=1)
doc.add_heading('8.1 数据集验收', level=2)
add_para(
    '数据集验收从数量、质量、分布、格式四个维度进行。'
    '数量：每类正例≥500，负例≥100，总计标注样本≥8,520，标准答案 JSON 文件数量与标注图像数量一致。'
    '质量：随机抽 200 张人工核查，标注准确率≥90%。交叉验证 IoU≥0.85，标签一致率≥95%。'
    '分布：光照/天气/角度各维度覆盖率≥10%。'
    '格式：YOLO .txt 格式无误，class_id 在 0-11 范围，坐标归一化 0-1。')
add_para(
    '标准答案 JSON 验收（本次新增）：JSON Schema 验证通过率 = 100%（自动化检查，一条不过即整批退回）；'
    '必填字段非空率 = 100%（image、bbox、class_id、inspection_item、violation_description、'
    'regulation_clause、regulation_text 全部非空）；'
    '法规条款引用准确率 ≥ 95%（随机抽 200 条人工核查，regulation_clause 和 regulation_text 与隐患类别匹配）；'
    '违规描述有效性检查：随机抽 100 条违规描述，判定为"有意义且与图像内容相关"的比例 ≥ 90%，'
    '同一类别的违规描述应涵盖至少 3 种不同的真实违规场景表述（避免千篇一律的复制粘贴）。')

doc.add_heading('8.2 模型验收', level=2)
add_para(
    'YOLO mAP@0.5 ≥92%，Recall ≥93%，Precision ≥85%。VLM 复核准确率 ≥90%。'
    '任意单类别指标退化 ≤2%。灰度期人工复核确认准确率 ≥85%。')
add_para(
    'VLM JSON 输出验收（本次新增）：VLM 推理输出的 JSON 格式合规率（能通过 Schema 验证的比例）≥ 95%；'
    '推理输出的 regulation_clause 字段与实际隐患类别匹配率 ≥ 90%。')

doc.add_heading('8.3 交付物清单', level=2)
add_bullet('《AI模型调优方案》（本文档）— PDF + 可编辑版')
add_bullet('《标注指导手册》— 含 12 类完整判定标准、正例/负例/边界图示、JSON 标准答案字段填写指南')
add_bullet('描述模板库 — 每类 3-5 个违规描述标准模板')
add_bullet('清洗后的原始图像数据集 — JPG 图片集')
add_bullet('标注完成的数据集 — JPG + YOLO .txt，含 train/val/test 划分')
add_bullet('标准答案 JSON 数据集 — 每条包含图片引用 + 人工标记 + 隐患内容（检查项目/违规描述/规定条款/规定原文），JSON Lines 格式，含计数型特殊字段')
add_bullet('隐患库 JSON 文件 — hazard_knowledge_base.json，含定义+法规引用+预填字段')
add_bullet('微调后 YOLO 权重文件 — .pt 格式')
add_bullet('微调后 VLM LoRA 适配器 — .safetensors 或 .bin 格式')
add_bullet('模型评估报告 — 新旧版本逐类对比，含 VLM JSON Schema 合规率统计')
add_bullet('灰度测试报告 + 项目总结报告')

doc.add_page_break()

# ═══════════════ 九、附录 ═══════════════
doc.add_heading('九、附录', level=1)
doc.add_heading('附录 A：12 类隐患 class_id 与检测类型对照表', level=2)
add_simple_table(
    ['class_id', '中文名称', '英文名称', '类别', '检测类型'],
    [
        ['0', '安全帽佩戴识别', 'no_hard_hat', '行为', '有/无'],
        ['1', '安全绳佩戴识别', 'no_safety_harness', '行为', '有/无'],
        ['2', '临边防护缺失识别', 'edge_protection', '物理设施', '有/无'],
        ['3', '施工电梯搭载人数识别', 'elevator_overload', '行为', '计数'],
        ['4', '工人私自开机识别', 'unauthorized_operation', '行为', '有/无'],
        ['5', '坑边堆载识别', 'pit_overload', '物理设施', '有/无'],
        ['6', '违规抽烟识别', 'smoking', '行为', '有/无'],
        ['7', '散料吊装识别', 'bulk_lifting', '物理设施', '有/无'],
        ['8', '吊篮作业人数识别', 'platform_overload', '物理设施', '计数'],
        ['9', '起重机械安拆人数识别', 'crane_assembly_count', '行为', '计数'],
        ['10', '出入口安全防护棚识别', 'safety_shelter', '物理设施', '有/无'],
        ['11', '工地围挡封闭识别', 'enclosure_breach', '物理设施', '有/无'],
    ]
)

doc.add_heading('附录 B：计数型隐患法规上限对照表', level=2)
add_simple_table(
    ['隐患', '法规上限', 'count_limit 值', '法规依据'],
    [
        ['#4 施工电梯搭载人数', '≤ 9 人（含司机）', '9', 'JGJ 33-2012 §2.0.7；韶建电[2017]68 号'],
        ['#9 吊篮作业人数', '≤ 2 人', '2', 'JGJ 202-2010 §5.5.8'],
        ['#9 双动力吊篮', '不允许单独 1 人（需 2 人）', '2（下限1）', 'JB/T 11699-2013 §5.2.3.7'],
        ['#10 起重机械安拆人数', '按专项施工方案要求', '待确认', '需客户提供具体方案人数'],
    ]
)

doc.add_heading('附录 C：术语表', level=2)
add_para('YOLO（You Only Look Once）：实时目标检测算法，将目标检测任务建模为单次回归问题，在速度和精度之间取得了良好平衡。')
add_para('VLM（Vision Language Model）：视觉语言大模型，能同时理解图像内容和自然语言描述。当前使用 Qwen3.6-27B 版本。')
add_para('LoRA（Low-Rank Adaptation）：参数高效微调方法，在原模型旁路添加低秩矩阵，不修改原始权重，适配器仅几十 MB。')
add_para('mAP（mean Average Precision）：目标检测综合评估指标，mAP@0.5 表示以 IoU≥0.5 作为正确检测的判定门限。')
add_para('IoU（Intersection over Union，交并比）：衡量预测边界框与真实边界框重叠程度的指标，取值 0-1。')
add_para('边界样本（Borderline Sample）：无法明确判定是否存在隐患的模糊场景，标注为 -1，不参与训练但作为难例分析资料。')
add_para('计数型隐患：指需要 AI 统计场景中人员数量的隐患类型（#4 施工电梯、#9 吊篮、#10 起重机械安拆），与传统的"有/无"判断相对。')
add_para('标准答案 JSON（Ground Truth JSON）：为每张训练图像配套编写的结构化标注文件，包含图像引用、边界框标注和隐患内容（检查项目/违规描述/规定条款/规定原文），是 VLM 监督微调的训练目标。')

# ── NEW: 附录 D 标准答案JSON Schema完整定义 ──
doc.add_heading('附录 D：标准答案 JSON Schema 完整定义', level=2)
add_para(
    '以下为所有参与方提供一份统一的 JSON Schema 参考。该 Schema 适用于标注工具的导出校验'
    '和数据集入库前的自动化验证。标注工具和复核脚本均以此 Schema 为准。')

schema_text = '''{
  "$schema": "https://json-schema.org/draft/2020-12/schema",
  "title": "标准答案JSON",
  "description": "AI隐患识别模型训练用标准答案数据格式",
  "type": "object",
  "required": ["dataset_version", "image", "image_info", "annotations"],
  "properties": {
    "dataset_version": {
      "type": "string",
      "description": "数据集版本号，如 1.0"
    },
    "image": {
      "type": "object",
      "required": ["file_name", "width", "height"],
      "properties": {
        "file_name": {"type": "string"},
        "width": {"type": "integer", "minimum": 640},
        "height": {"type": "integer", "minimum": 480}
      }
    },
    "image_info": {
      "type": "object",
      "required": ["camera_id", "capture_time", "light_condition", "weather", "construction_phase"],
      "properties": {
        "camera_id": {"type": "string"},
        "capture_time": {"type": "string", "format": "date-time"},
        "light_condition": {"type": "string", "enum": ["daylight","overcast","dusk","night","backlight"]},
        "weather": {"type": "string", "enum": ["sunny","cloudy","rainy"]},
        "construction_phase": {"type": "string", "enum": ["foundation","main_structure","decoration"]}
      }
    },
    "annotations": {
      "type": "array",
      "minItems": 1,
      "items": {
        "type": "object",
        "required": ["annotation_id","bbox","class_id","class_name","detection_type",
                     "is_positive","is_borderline","annotator_id","annotation_date","hazard_content"],
        "properties": {
          "annotation_id": {"type": "integer", "minimum": 1},
          "bbox": {
            "type": "object",
            "required": ["x","y","width","height"],
            "properties": {
              "x": {"type": "integer", "minimum": 0},
              "y": {"type": "integer", "minimum": 0},
              "width": {"type": "integer", "minimum": 1},
              "height": {"type": "integer", "minimum": 1}
            }
          },
          "class_id": {"type": "integer", "minimum": 0, "maximum": 11},
          "class_name": {"type": "string"},
          "detection_type": {"type": "string", "enum": ["有/无","计数"]},
          "is_positive": {"type": "boolean"},
          "is_borderline": {"type": "boolean"},
          "borderline_reason": {"type": "string"},
          "count": {"type": ["integer", "null"]},
          "count_limit": {"type": ["integer", "null"]},
          "annotator_id": {"type": "string", "pattern": "^A[0-9]{3}$"},
          "annotation_date": {"type": "string", "format": "date"},
          "hazard_content": {
            "type": "object",
            "required": ["inspection_item","violation_description","regulation_clause",
                         "regulation_text","kb_entry_id","severity"],
            "properties": {
              "inspection_item": {"type": "string"},
              "violation_description": {"type": "string", "minLength": 10},
              "regulation_clause": {"type": "string"},
              "regulation_text": {"type": "string", "minLength": 5},
              "kb_entry_id": {"type": "string", "pattern": "^KB-HZ-[0-9]{3}$"},
              "severity": {"type": "string", "enum": ["高","中","低"]}
            }
          }
        }
      }
    }
  },
  "if": {
    "properties": {"detection_type": {"const": "计数"}}
  },
  "then": {
    "required": ["count", "count_limit"]
  }
}'''
add_code_block(schema_text)

add_para(
    'Schema 使用说明：所有 JSON 文件入库前须通过上述 Schema 的自动化验证。关键校验规则包括：'
    'class_id 必须在 0-11 范围内；detection_type 为"计数"时 count 和 count_limit 为必填字段；'
    'violation_description 至少 10 个字符（禁止空白或无意义填充）；'
    'annotator_id 格式为 A 开头加三位数字（如 A001）；kb_entry_id 格式为 KB-HZ-XXX；'
    'severity 取值仅限"高""中""低"。')

doc.add_page_break()

# ── NEW: 附录 E 标准答案JSON示例 ──
doc.add_heading('附录 E：标准答案 JSON 示例（典型类别）', level=2)
add_para(
    '以下选取 3 个典型类别给出完整的标准答案 JSON 示例，涵盖"有/无型""计数型""物理设施型"三种场景，'
    '供标注人员、复核人员和 AI 工程师参考。')

doc.add_heading('E.1 示例一：安全帽佩戴识别（#1，有/无型）', level=3)
add_code_block('''{
  "dataset_version": "1.0",
  "image": {
    "file_name": "CAM01_20250701_083000.jpg",
    "width": 1920, "height": 1080
  },
  "image_info": {
    "camera_id": "CAM-01",
    "capture_time": "2025-07-01T08:30:00",
    "light_condition": "daylight",
    "weather": "sunny",
    "construction_phase": "main_structure"
  },
  "annotations": [
    {
      "annotation_id": 1,
      "bbox": {"x": 450, "y": 300, "width": 80, "height": 100},
      "class_id": 0,
      "class_name": "安全帽佩戴识别",
      "detection_type": "有/无",
      "is_positive": true,
      "is_borderline": false,
      "borderline_reason": "",
      "count": null,
      "count_limit": null,
      "annotator_id": "A001",
      "annotation_date": "2025-07-05",
      "hazard_content": {
        "inspection_item": "安全帽佩戴识别",
        "violation_description": "主体结构施工层一名作业人员头部未见安全帽，佩戴普通棒球帽替代",
        "regulation_clause": "JGJ 184-2009 §2.0.4",
        "regulation_text": "进入施工现场人员必须佩戴安全帽。",
        "kb_entry_id": "KB-HZ-001",
        "severity": "中"
      }
    }
  ]
}''')

doc.add_heading('E.2 示例二：施工电梯搭载人数识别（#4，计数型）', level=3)
add_code_block('''{
  "dataset_version": "1.0",
  "image": {
    "file_name": "CAM05_20250702_170500.jpg",
    "width": 1920, "height": 1080
  },
  "image_info": {
    "camera_id": "CAM-05",
    "capture_time": "2025-07-02T17:05:00",
    "light_condition": "dusk",
    "weather": "cloudy",
    "construction_phase": "main_structure"
  },
  "annotations": [
    {
      "annotation_id": 1,
      "bbox": {"x": 120, "y": 200, "width": 350, "height": 420},
      "class_id": 3,
      "class_name": "施工电梯搭载人数识别",
      "detection_type": "计数",
      "is_positive": true,
      "is_borderline": false,
      "borderline_reason": "",
      "count": 11,
      "count_limit": 9,
      "annotator_id": "A003",
      "annotation_date": "2025-07-06",
      "hazard_content": {
        "inspection_item": "施工电梯搭载人数识别",
        "violation_description": "施工电梯轿厢内可见11人（含司机），超过法规上限9人，属于严重超载",
        "regulation_clause": "JGJ 33-2012 §2.0.7; 韶建电[2017]68号",
        "regulation_text": "承载人数不得超过9人（含司机）。",
        "kb_entry_id": "KB-HZ-004",
        "severity": "高"
      }
    }
  ]
}''')

doc.add_heading('E.3 示例三：坑边堆载识别（#6，物理设施型，含多隐患）', level=3)
add_code_block('''{
  "dataset_version": "1.0",
  "image": {
    "file_name": "CAM02_20250703_141500.jpg",
    "width": 1920, "height": 1080
  },
  "image_info": {
    "camera_id": "CAM-02",
    "capture_time": "2025-07-03T14:15:00",
    "light_condition": "daylight",
    "weather": "sunny",
    "construction_phase": "foundation"
  },
  "annotations": [
    {
      "annotation_id": 1,
      "bbox": {"x": 200, "y": 350, "width": 500, "height": 300},
      "class_id": 5,
      "class_name": "坑边堆载识别",
      "detection_type": "有/无",
      "is_positive": true,
      "is_borderline": false,
      "borderline_reason": "",
      "count": null,
      "count_limit": null,
      "annotator_id": "A002",
      "annotation_date": "2025-07-06",
      "hazard_content": {
        "inspection_item": "坑边堆载识别",
        "violation_description": "基坑边缘约0.8m处堆放大量钢筋和模板材料，堆载距离明显小于1.5m安全间距，且基坑周边地面未做硬化处理",
        "regulation_clause": "JGJ 120-2012 §8.1.5; JGJ 311-2013 §11.2.2",
        "regulation_text": "基坑周边施工材料、设施或车辆荷载严禁超过设计要求的地面荷载限值。基坑周边1.5m范围内不宜堆载。",
        "kb_entry_id": "KB-HZ-006",
        "severity": "高"
      }
    },
    {
      "annotation_id": 2,
      "bbox": {"x": 750, "y": 400, "width": 90, "height": 110},
      "class_id": 0,
      "class_name": "安全帽佩戴识别",
      "detection_type": "有/无",
      "is_positive": true,
      "is_borderline": false,
      "borderline_reason": "",
      "count": null,
      "count_limit": null,
      "annotator_id": "A002",
      "annotation_date": "2025-07-06",
      "hazard_content": {
        "inspection_item": "安全帽佩戴识别",
        "violation_description": "基坑边一名管理人员未佩戴安全帽",
        "regulation_clause": "JGJ 184-2009 §2.0.4",
        "regulation_text": "进入施工现场人员必须佩戴安全帽。",
        "kb_entry_id": "KB-HZ-001",
        "severity": "中"
      }
    }
  ]
}''')

add_para(
    '示例三展示了一张图像中包含两个不同类别隐患的场景（坑边堆载 + 安全帽缺失）。'
    '标注时每个隐患独立绘制边界框、选择类别、填写 hazard_content，'
    'annotations 数组包含两个元素，各自携带完整的证据链信息。')

# ── Save ──────────────────────────────────────────────
output_path = r'f:\0.AI设计库\ai视频识别\现版本设计文件\AI隐患识别模型调优方案_V2.2.docx'
doc.save(output_path)
print(f'[OK] Saved: {output_path}')
